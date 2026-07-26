from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import os

from docx import Document
from docx.text.paragraph import Paragraph


DOCX_PATH = Path(__file__).with_name("ApiStation-FinOps-产品与技术设计-v0.1.docx")


def find_paragraph(document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise ValueError(f"paragraph not found: {text}")


def maybe_find_paragraph(document, text: str) -> Paragraph | None:
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    return None


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_paragraph(document, old: str, new: str) -> None:
    paragraph = maybe_find_paragraph(document, old)
    if paragraph is not None:
        set_paragraph_text(paragraph, new)
    elif maybe_find_paragraph(document, new) is None:
        raise ValueError(f"paragraph not found for replacement: {old}")


def copy_numbering(source: Paragraph, target: Paragraph) -> None:
    if source._p.pPr is None or source._p.pPr.numPr is None:
        return
    target_p_pr = target._p.get_or_add_pPr()
    if target_p_pr.numPr is not None:
        target_p_pr.remove(target_p_pr.numPr)
    target_p_pr.append(deepcopy(source._p.pPr.numPr))


def insert_after(reference: Paragraph, text: str, style: str | None = None, numbering_source: Paragraph | None = None) -> Paragraph:
    paragraph = reference._parent.add_paragraph(text, style=style)
    reference._p.addnext(paragraph._p)
    if numbering_source is not None:
        copy_numbering(numbering_source, paragraph)
    return paragraph


def replace_in_tables(document, old: str, new: str) -> None:
    found = False
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        set_paragraph_text(paragraph, paragraph.text.replace(old, new))
                        found = True
    if not found and not any(new in cell.text for table in document.tables for row in table.rows for cell in row.cells):
        raise ValueError(f"table text not found for replacement: {old}")


def main() -> None:
    document = Document(DOCX_PATH)

    replacements = {
        "6.2 用户利润分析": "6.2 用户账务与利润",
        "6.3 账号成本中心": "6.3 账号台账与成本",
        "6.4 用量分析": "6.5 用量与扣费",
        "6.5 充值与资金": "6.6 充值与资金",
        "6.6 成本规则": "6.7 成本核算",
        "6.7 对账中心": "6.8 对账中心",
        "6.8 报表、导出与告警": "6.10 告警、报表与导出",
        "经营总览、用户利润、账号成本、用量分析、充值与资金、对账中心、成本规则 7 个页面；":
            "经营总览、用户账务与利润、用量与扣费、账号台账与成本、供应商与采购、成本核算、充值与资金、对账中心、数据同步、告警中心 10 个页面；",
        "分组左侧菜单、右侧数据工作区、时间范围、搜索、CSV 导出、演示/数据库状态，以及成本模板、账号采购成本和手工支出录入交互；":
            "四组固定左侧菜单、右侧数据工作区、时间范围、搜索、CSV 导出、演示/数据库状态，以及成本模板、账号采购成本、供应商聚合、手工支出和来源级同步状态；",
        "apistation-finops/web/：经营总览、用户利润、账号成本、用量、资金、对账和成本规则页面实现。":
            "apistation-finops/web/：经营、用户、用量、账号、供应商采购、成本核算、资金、对账、同步和告警页面实现。",
    }
    for old, new in replacements.items():
        replace_paragraph(document, old, new)

    if maybe_find_paragraph(document, "6.4 供应商与采购") is None:
        anchor = find_paragraph(document, "支持账号标签：供应商、采购渠道、批次、负责人、地区、质量等级、是否共享、成本中心、自定义标签。标签必须保存历史快照，避免改标签后历史报表整体漂移。")
        anchor = insert_after(anchor, "6.4 供应商与采购", "Heading 2")
        anchor = insert_after(anchor, "供应商页面按上游归集账号、平台、请求、Token、确认收入、已入账人民币成本、期间采购、成本覆盖状态和已入账毛利，并明确显示成本待补、未标记供应商和即将到期账号。", "Normal")
        insert_after(anchor, "采购批次台账至少包括账号、供应商、批次、成本模板、人民币成本、手续费、税费、生效期和状态。采购成本与现金支出需要建立稳定关联；首版允许双录并在对账中心暴露未关联差异，后续改为一次提交同时生成成本期间和现金流水。", "Normal")

    if maybe_find_paragraph(document, "6.9 数据同步") is None:
        report_heading = find_paragraph(document, "6.10 告警、报表与导出")
        bullet_source = find_paragraph(document, "日报、周报、月报：现金、消耗、成本、毛利、用量和异常摘要；")
        sync_heading = report_heading.insert_paragraph_before("6.9 数据同步", style="Heading 2")
        for text in [
            "按用户与账号、用量与扣费、充值与退款、自动对账四类来源展示状态；",
            "展示最近成功时间、当前游标、最大延迟、累计行数和最近错误；",
            "同步失败必须写入状态台账，不能只打印服务日志；",
            "后续增加手工重跑、历史回填、/ready 和 Schema 兼容结果。",
        ]:
            paragraph = report_heading.insert_paragraph_before(text, style="Normal")
            copy_numbering(bullet_source, paragraph)

    replace_in_tables(
        document,
        "已有 7 个管理页面、演示/数据库模式和 CSV；待详情下钻、完整台账与高级筛选",
        "已有 10 个管理页面、演示/数据库模式和 CSV；待请求明细下钻、充值页签、完整台账与高级筛选",
    )

    forbidden = ["用户利润分析", "账号成本中心", "6.4 用量分析", "7 个页面", "已有 7 个管理页面"]
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    full_text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    remaining = [text for text in forbidden if text in full_text]
    if remaining:
        raise ValueError(f"stale navigation text remains: {remaining}")

    document.core_properties.modified = datetime(2026, 7, 26, 0, 0, 0)
    temp_path = DOCX_PATH.with_suffix(".tmp.docx")
    document.save(temp_path)
    os.replace(temp_path, DOCX_PATH)


if __name__ == "__main__":
    main()
