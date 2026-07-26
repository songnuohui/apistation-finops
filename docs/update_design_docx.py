from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import os

from docx import Document


DOCX_PATH = Path(__file__).with_name("ApiStation-FinOps-产品与技术设计-v0.1.docx")


def find_paragraph(document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise ValueError(f"paragraph not found: {text}")


def maybe_find_paragraph(document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    return None


def replace_paragraph(document, old: str, new: str) -> None:
    paragraph = maybe_find_paragraph(document, old)
    if paragraph is None:
        if maybe_find_paragraph(document, new) is None:
            raise ValueError(f"paragraph not found for replacement: {old}")
        return
    if not paragraph.runs:
        paragraph.add_run(new)
        return
    paragraph.runs[0].text = new
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_in_runs(document, paragraph_text: str, replacements: dict[str, str]) -> None:
    paragraph = find_paragraph(document, paragraph_text)
    for run in paragraph.runs:
        for old, new in replacements.items():
            run.text = run.text.replace(old, new)


def copy_numbering(source, target) -> None:
    source_num_pr = source._p.pPr.numPr
    target_p_pr = target._p.get_or_add_pPr()
    if target_p_pr.numPr is not None:
        target_p_pr.remove(target_p_pr.numPr)
    target_p_pr.append(deepcopy(source_num_pr))


def insert_before(reference, text: str, style: str | None = None, numbering_source=None):
    paragraph = reference.insert_paragraph_before(text, style=style)
    if numbering_source is not None:
        copy_numbering(numbering_source, paragraph)
    return paragraph


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def main() -> None:
    document = Document(DOCX_PATH)

    version = find_paragraph(document, "文档版本：v0.1（讨论稿）")
    version.runs[-1].text = "v0.1（实施同步稿）"
    date = find_paragraph(document, "日期：2026-07-15")
    date.runs[-1].text = "2026-07-26"

    replace_in_runs(
        document,
        "有条件可行。 当前规模较小，新增一个轻量 Go 管理服务不会成为主要负担；30M 带宽对统计元数据足够。主要风险是 2 核 CPU 与 4GB 内存下，统计查询和 ApiStation 转发争抢 PostgreSQL、CPU 和磁盘 I/O。",
        {"轻量 Go 管理服务": "轻量 Node.js 管理服务"},
    )
    replace_paragraph(
        document,
        "FinOps 应用为一个 Go 容器，前端静态文件内嵌；",
        "FinOps 应用为一个 Node.js 单服务容器，前端静态文件由同一进程托管；",
    )
    replace_in_runs(
        document,
        "所有货币字段在 PostgreSQL 使用 NUMERIC，Go 端使用十进制定点库，不使用 float64 做最终台账计算。",
        {
            "Go 端使用十进制定点库": "Node.js 端使用 decimal.js 做十进制定点计算",
            "float64": "JavaScript Number",
        },
    )

    replace_paragraph(
        document,
        "后端：Go，与 ApiStation 团队技能和部署方式一致；",
        "后端：Node.js 22 ESM，使用内置 HTTP 服务承载 API 和静态文件，pg 连接 PostgreSQL，decimal.js 处理金额；",
    )
    replace_paragraph(
        document,
        "前端：Vue 3 + TypeScript + Vite，沿用现有界面交互习惯；",
        "前端：原生 HTML、CSS 和 JavaScript，由同一 Node.js 服务托管，沿用 ApiStation 左侧菜单与右侧数据工作区的管理端布局；",
    )
    replace_paragraph(
        document,
        "图表：Chart.js；",
        "图表：浏览器原生 Canvas，首版不额外引入图表框架；",
    )
    replace_paragraph(
        document,
        "部署：Docker Compose，单个 FinOps 镜像内嵌前端；",
        "部署目标：Docker Compose，单个 FinOps 镜像同时运行 API、同步器和静态管理前端；",
    )
    replace_in_runs(
        document,
        "apistation-finops：新增 Go 应用，内部端口 8090；",
        {"新增 Go 应用": "新增 Node.js 单服务应用"},
    )
    replace_in_runs(
        document,
        "建议参数：FinOps GOMAXPROCS=1 或 2、数据库最大连接 5、空闲连接 2、同步批次 1,000、同步周期 60 秒、查询超时 5 秒、后台重算单并发、容器内存上限 384MB。",
        {
            "GOMAXPROCS=1": "NODE_OPTIONS=--max-old-space-size=256",
            " 或 2、": "、",
        },
    )
    replace_paragraph(document, "8.5 可观测性与运维", "8.5 可观测性与运维目标")
    replace_paragraph(
        document,
        "/health：进程和数据库连接；",
        "/health（已实现）：返回进程状态、演示/数据库模式和运行时间；",
    )
    replace_paragraph(
        document,
        "/ready：Schema 版本、来源权限、最近同步成功；",
        "/ready（待实现）：Schema 版本、来源权限、数据库连接和最近同步成功；",
    )

    replace_paragraph(document, "11. 实施计划与工作量", "11. 实施状态与工作计划")
    plan_intro = find_paragraph(
        document,
        "以一名熟悉现有代码的全栈开发者为参考，MVP 建议按 4 个阶段推进，整体约 3-5 周；若上游账单格式复杂或需要改 ApiStation 事件链路，时间需增加。",
    )
    if maybe_find_paragraph(document, "11.1 当前实施状态（截至 2026-07-26）") is None:
        bullet_template = find_paragraph(document, "FinOps 不部署第二套 PostgreSQL 和 Redis；")
        insert_before(plan_intro, "11.1 当前实施状态（截至 2026-07-26）", "Heading 2")
        insert_before(
            plan_intro,
            "当前已从方案评审进入 MVP 原型开发，技术路线已冻结为 Node.js 单服务。代码基线已具备：",
            "Normal",
        )
        implemented = [
            "Node.js 22 ESM 单进程同时承载 HTTP API、周期同步器和静态管理前端；未配置 DATABASE_URL 时自动进入演示模式；",
            "PostgreSQL finops Schema 初始迁移，覆盖同步游标、用户/账号维度、请求事实、日聚合、成本档案、采购期间、现金流水、对账和审计；",
            "用户、账号、usage_logs 和支付订单增量同步，包含事务、游标、幂等写入、日聚合和近期用量核对；",
            "经营总览、用户账务与利润、用量与扣费、账号台账与成本、供应商与采购、成本核算、充值与资金、对账中心、数据同步、告警中心 10 个页面；",
            "四组固定左侧菜单、右侧数据工作区、时间范围、搜索、CSV 导出、演示/数据库状态，以及成本模板、账号采购成本、供应商聚合、手工支出和来源级同步状态；",
            "金额换算与固定成本分摊基础函数单元测试。",
            "Dockerfile、Compose 示例、最小数据库权限脚本、资源限制和部署说明。",
        ]
        for item in implemented:
            insert_before(plan_intro, item, "Normal", bullet_template)
        insert_before(plan_intro, "当前尚未达到生产上线标准，剩余重点包括：", "Normal")
        pending = [
            "使用真实 ApiStation PostgreSQL 完成 Schema 兼容、同步、回填、金额与 Token 对账；",
            "补齐退款、赠送、返利、余额流水、固定成本分摊结果和冲正台账等完整闭环；",
            "完成真实环境鉴权联调、告警、/ready、备份恢复、压测和连续 7 天影子运行。",
        ]
        for item in pending:
            insert_before(plan_intro, item, "Normal", bullet_template)
        insert_before(plan_intro, "11.2 剩余实施计划", "Heading 2")

    replace_paragraph(
        document,
        "以一名熟悉现有代码的全栈开发者为参考，MVP 建议按 4 个阶段推进，整体约 3-5 周；若上游账单格式复杂或需要改 ApiStation 事件链路，时间需增加。",
        "原始 MVP 参考工作量为 3-5 周。当前已完成口径冻结、数据底座和核心页面的第一版代码，后续时间主要取决于真实数据兼容、历史账务完整性和上游账单格式。",
    )

    plan_table = document.tables[8]
    plan_rows = [
        ["阶段", "当前状态与剩余交付", "参考时间"],
        ["0. 口径冻结", "已完成：本位币、余额单位、分摊、保留期和访问范围按默认值冻结", "已完成"],
        ["1. 数据底座", "已有 Schema、游标同步、幂等和日聚合；待真实库验证、回填与同步监控", "3-5 天"],
        ["2. 核心产品", "已有 10 个管理页面、演示/数据库模式和 CSV；待请求明细下钻、充值页签、完整台账与高级筛选", "5-8 天"],
        ["3. 对账上线", "待完成生产级对账、告警、权限、审计、压测、部署与备份", "5-8 天"],
    ]
    for row, values in zip(plan_table.rows, plan_rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)

    replace_paragraph(
        document,
        "第一周不要先做漂亮大屏，应先用真实数据跑通 5 个校验样本：一个在线充值用户、一个兑换码用户、一个有返利用户、一个固定订阅账号、一个按量账号。",
        "进入真实数据库联调后，优先跑通 5 个校验样本：一个在线充值用户、一个兑换码用户、一个有返利用户、一个固定订阅账号、一个按量账号。",
    )

    replace_paragraph(document, "12. 待确认事项与建议默认值", "12. 已冻结默认值与后续可配置项")
    defaults_table = document.tables[9]
    set_cell_text(defaults_table.cell(0, 0), "默认项")
    set_cell_text(defaults_table.cell(0, 1), "当前取值")
    replace_paragraph(
        document,
        "需要您优先拍板的只有四项：本位币、固定账号成本分摊方法、历史数据保留期限、是否能取得上游真实账单。其余可以先按建议默认值实施。",
        "首版开发已按上表默认值执行，不再作为开发阻塞项。后续接入上游真实人民币账单或开放更多角色时，再通过版本化配置调整，不回写污染历史结果。",
    )

    replace_paragraph(
        document,
        "本方案基于当前工作区源码和 7 张系统截图，关键依据包括：",
        "本方案基于当前工作区源码、首轮 7 张功能截图和追加 3 张管理端布局参考，关键依据包括：",
    )
    evidence_anchor = find_paragraph(
        document,
        "当前未获得服务器磁盘、实际内存峰值、PostgreSQL 表行数、支付渠道手续费结算规则和上游账单样例，因此资源预算与自动对账范围仍需在实施前用真实环境数据校准。",
    )
    evidence_items = [
        "apistation-finops/src/server.mjs 与 src/services/sync-service.mjs：Node.js 单服务、API、静态前端托管和增量同步实现；",
        "apistation-finops/migrations/001_init.sql：当前 finops Schema 初始表结构；",
        "apistation-finops/web/：经营、用户、用量、账号、供应商采购、成本核算、资金、对账、同步和告警页面实现。",
    ]
    if maybe_find_paragraph(document, evidence_items[0]) is None:
        evidence_template = find_paragraph(document, "deploy/docker-compose.yml：现有单应用、PostgreSQL、Redis 和 Docker 网络部署方式。")
        for item in evidence_items:
            insert_before(evidence_anchor, item, "Normal", evidence_template)

    replace_paragraph(document, "14. 推荐的下一次评审议程", "14. 开发期下一次评审议程")
    review_replacements = [
        (
            "用 3 笔真实充值和 5 条真实请求逐项确认金额含义；",
            "按追加截图复核左侧菜单、右侧数据工作区、筛选条、汇总卡和明细表的桌面/移动端一致性；",
        ),
        (
            "选取 2 个固定订阅账号和 1 个按量账号，确认采购成本及分摊；",
            "使用真实 ApiStation Schema 跑通用户、账号、usage_logs 和支付订单的只读增量同步；",
        ),
        (
            "确认人民币本位币、赠送与返利处理方式；",
            "用 3 笔真实充值、5 条真实请求、2 个固定订阅账号和 1 个按量账号完成金额守恒与成本分摊校验；",
        ),
        (
            "确认首版页面优先级和谁可以访问；",
            "检查服务器磁盘、容器内存峰值、usage_logs 行数与日增量，并完成资源限额压测；",
        ),
        (
            "查看服务器磁盘、容器内存峰值、usage_logs 行数与日增量；",
            "完成 Docker 部署、数据库最小权限、备份恢复后，进入连续 7 天影子运行。",
        ),
    ]
    for old, new in review_replacements:
        replace_paragraph(document, old, new)
    obsolete_review = maybe_find_paragraph(document, "冻结 v0.2 指标字典后，再进入数据库设计与交互原型。")
    if obsolete_review is not None:
        remove_paragraph(obsolete_review)

    document.core_properties.modified = datetime(2026, 7, 26, 0, 0, 0)
    temp_path = DOCX_PATH.with_suffix(".tmp.docx")
    document.save(temp_path)
    os.replace(temp_path, DOCX_PATH)


if __name__ == "__main__":
    main()
